"""
Document Loader — reads .txt, .md, .pdf, .docx, .csv and returns raw text + metadata.
"""
from __future__ import annotations

import io
import os
from dataclasses import dataclass
from pathlib import Path
from typing import Optional

from loguru import logger


@dataclass
class LoadedDocument:
    filename: str
    file_size: int          # bytes
    raw_text: str
    mime_type: str
    page_count: Optional[int] = None   # PDFs only
    row_count: Optional[int] = None    # CSVs only


SUPPORTED_EXTENSIONS = {".txt", ".md", ".pdf", ".docx", ".csv", ".tsv", ".xlsx"}


def load_document(file_bytes: bytes, filename: str) -> LoadedDocument:
    """
    Load a document from raw bytes and return a LoadedDocument.
    Dispatches to the appropriate parser based on file extension.
    """
    ext = Path(filename).suffix.lower()
    file_size = len(file_bytes)

    if ext not in SUPPORTED_EXTENSIONS:
        raise ValueError(f"Unsupported file type: {ext}. Supported: {SUPPORTED_EXTENSIONS}")

    logger.info(f"Loading document: {filename} ({file_size:,} bytes, type={ext})")

    if ext in (".txt", ".md"):
        text = _load_text(file_bytes, filename)
        return LoadedDocument(filename=filename, file_size=file_size, raw_text=text, mime_type="text/plain")

    elif ext == ".pdf":
        text, page_count = _load_pdf(file_bytes, filename)
        return LoadedDocument(
            filename=filename, file_size=file_size, raw_text=text,
            mime_type="application/pdf", page_count=page_count
        )

    elif ext == ".docx":
        text = _load_docx(file_bytes, filename)
        return LoadedDocument(filename=filename, file_size=file_size, raw_text=text,
                              mime_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

    elif ext in (".csv", ".tsv", ".xlsx"):
        text, row_count = _load_tabular(file_bytes, filename, ext)
        return LoadedDocument(filename=filename, file_size=file_size, raw_text=text,
                              mime_type="text/csv", row_count=row_count)

    raise ValueError(f"Unhandled extension: {ext}")


def _load_text(file_bytes: bytes, filename: str) -> str:
    """Decode text files, trying UTF-8 first then latin-1."""
    for encoding in ("utf-8", "utf-8-sig", "latin-1"):
        try:
            return file_bytes.decode(encoding)
        except UnicodeDecodeError:
            continue
    raise ValueError(f"Could not decode {filename} with any known encoding")


def _load_pdf(file_bytes: bytes, filename: str) -> tuple[str, int]:
    """Extract text from PDF using PyMuPDF (fitz)."""
    try:
        import fitz  # PyMuPDF
    except ImportError:
        raise ImportError("PyMuPDF not installed. Run: pip install PyMuPDF")

    doc = fitz.open(stream=file_bytes, filetype="pdf")
    pages = []
    for page_num, page in enumerate(doc, start=1):
        page_text = page.get_text("text")
        if page_text.strip():
            pages.append(f"[Page {page_num}]\n{page_text}")

    doc.close()
    full_text = "\n\n".join(pages)
    logger.debug(f"PDF: {len(pages)} pages extracted from {filename}")
    return full_text, len(pages)


def _load_docx(file_bytes: bytes, filename: str) -> str:
    """Extract text from DOCX preserving paragraph structure."""
    try:
        from docx import Document
    except ImportError:
        raise ImportError("python-docx not installed. Run: pip install python-docx")

    doc = Document(io.BytesIO(file_bytes))
    paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
    # Also extract table content
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                paragraphs.append(row_text)

    return "\n\n".join(paragraphs)


def _load_tabular(file_bytes: bytes, filename: str, ext: str) -> tuple[str, int]:
    """Convert CSV/TSV/XLSX to a readable text representation."""
    try:
        import pandas as pd
    except ImportError:
        raise ImportError("pandas not installed. Run: pip install pandas")

    buf = io.BytesIO(file_bytes)

    if ext == ".csv":
        # Try to detect delimiter
        sample = file_bytes[:2048].decode("utf-8", errors="ignore")
        delimiter = "\t" if sample.count("\t") > sample.count(",") else ","
        df = pd.read_csv(buf, delimiter=delimiter, on_bad_lines="skip")
    elif ext == ".tsv":
        df = pd.read_csv(buf, delimiter="\t", on_bad_lines="skip")
    elif ext == ".xlsx":
        df = pd.read_excel(buf)

    row_count = len(df)
    # Convert to markdown-style table for readability
    lines = [f"Columns: {', '.join(df.columns.astype(str))}"]
    lines.append(df.to_string(index=False, max_rows=1000))
    if row_count > 1000:
        lines.append(f"... ({row_count - 1000} more rows truncated)")

    return "\n".join(lines), row_count
